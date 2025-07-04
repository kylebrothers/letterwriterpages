def handle_no_call_page(page_name, form_data, uploaded_files_data, server_files_data, session_id):
    """Handle no-call pages that organize and display information without API calls"""
    try:
        logger.info(f"Processing no-call page: {page_name}")
        
        # Build organized content from available data
        content_sections = []
        
        # Add page header
        page_title = form_data.get('page_title', page_name.replace('-', ' ').title())
        content_sections.append(f"# {page_title}")
        content_sections.append("")
        
        # Add server reference files with rich data
        if server_files_data:
            content_sections.append("## Reference Materials")
            content_sections.append("")
            for file_name, file_data in server_files_data.items():
                content_sections.append(f"### {file_name.title()}")
                content_sections.append("")
                
                if isinstance(file_data, dict):
                    # Rich file data
                    content_sections.append(f"**File Type:** {file_data.get('file_type', 'unknown').upper()}")
                    
                    if 'text_content' in file_data:
                        content_sections.append("**Text Content Preview:**")
                        preview = file_data['text_content'][:500] + "..." if len(file_data['text_content']) > 500 else file_data['text_content']
                        content_sections.append(preview)
                    
                    if 'xml_structure' in file_data and file_data['xml_structure']:
                        content_sections.append("**Document Structure Available:** Yes (XML)")
                    
                    if 'form_data' in file_data and file_data['form_data']:
                        content_sections.append("**PDF Form Data Available:** Yes")
                        if 'document_info' in file_data['form_data']:
                            doc_info = file_data['form_data']['document_info']
                            content_sections.append(f"- Pages: {doc_info.get('page_count', 'unknown')}")
                            content_sections.append(f"- Has Form Fields: {doc_info.get('has_form_fields', False)}")
                else:
                    # Legacy text data
                    preview = str(file_data)[:500] + "..." if len(str(file_data)) > 500 else str(file_data)
                    content_sections.append(preview)
                
                content_sections.append("")
        
        # Add form data
        if form_data:
            content_sections.append("## Submitted Information")
            content_sections.append("")
            for key, value in form_data.items():
                if key not in ['page_type', 'page_title'] and value.strip():
                    display_key = key.replace('_', ' ').title()
                    content_sections.append(f"**{display_key}:** {value}")
            content_sections.append("")
        
        # Add uploaded files with rich data
        if uploaded_files_data:
            content_sections.append("## Uploaded Documents")
            content_sections.append("")
            for file_type, file_data in uploaded_files_data.items():
                content_sections.append(f"### {file_type.title()}")
                content_sections.append("")
                
                if isinstance(file_data, dict):
                    # Rich file data
                    content_sections.append(f"**File Type:** {file_data.get('file_type', 'unknown').upper()}")
                    
                    if 'text_content' in file_data:
                        content_sections.append("**Text Content Preview:**")
                        preview = file_data['text_content'][:500] + "..." if len(file_data['text_content']) > 500 else file_data['text_content']
                        content_sections.append(preview)
                    
                    if 'xml_structure' in file_data and file_data['xml_structure']:
                        content_sections.append("**Document Structure:** Available")
                        xml_files = list(file_data['xml_structure'].keys())
                        if 'error' in xml_files:
                            xml_files.remove('error')
                        content_sections.append(f"- XML Components: {', '.join(xml_files)}")
                    
                    if 'form_data' in file_data and file_data['form_data']:
                        content_sections.append("**PDF Analysis:**")
                        if 'document_info' in file_data['form_data']:
                            doc_info = file_data['form_data']['document_info']
                            content_sections.append(f"- Pages: {doc_info.get('page_count', 'unknown')}")
                            content_sections.append(f"- Encrypted: {doc_info.get('is_encrypted', False)}")
                            content_sections.append(f"- Has Form Fields: {doc_info.get('has_form_fields', False)}")
                        
                        if 'metadata' in file_data['form_data'] and file_data['form_data']['metadata']:
                            content_sections.append("- Metadata: Available")
                else:
                    # Legacy text data
                    preview = str(file_data)[:500] + "..." if len(str(file_data)) > 500 else str(file_data)
                    content_sections.append(preview)
                
                content_sections.append("")
        
        # Add summary
        total_items = len(server_files_data) + len(uploaded_files_data) + len([k for k in form_data.keys() if k not in ['page_type', 'page_title'] and form_data[k].strip()])
        content_sections.append("## Summary")
        content_sections.append("")
        content_sections.append(f"- **Server Reference Files:** {len(server_files_data)}")
        content_sections.append(f"- **Uploaded Documents:** {len(uploaded_files_data)}")
        content_sections.append(f"- **Form Fields Completed:** {len([k for k in form_data.keys() if k not in ['page_type', 'page_title'] and form_data[k].strip()])}")
        content_sections.append(f"- **Total Items Processed:** {total_items}")
        
        generated_content = "\n".join(content_sections)
        
        logger.info(f"No-call page processed successfully for: {page_name}")
        
        return jsonify({
            'success': True,
            'content': generated_content,
            'session_id': session_id,
            'page_type': 'no-call',
            'files_processed': list(uploaded_files_data.keys()),
            'server_files_loaded': list(server_files_data.keys()),
            'form_fields_received': [k for k in form_data.keys() if k not in ['page_type', 'page_title']]
        })
        
    except Exception as e:
        logger.error(f"Error in no-call page handler: {e}")
        return jsonify({'error': 'Error processing no-call page'}), 500

def handle_claude_call_page(page_name, form_data, uploaded_files_data, server_files_data, session_id):
    """Handle Claude API pages (existing functionality)"""
    if not claude_client:
        return jsonify({'error': 'Claude API not available'}), 503
    
    try:
        # Build Claude prompt from form data, uploaded files, and server files
        claude_prompt, error = build_claude_prompt(form_data, uploaded_files_data, server_files_data)
        if error:
            return jsonify({'error': error}), 400
        
        # Call Claude API
        message = claude_client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=4000,
            messages=[
                {"role": "user", "content": claude_prompt}
            ]
        )
        
        generated_content = message.content[0].text
        
        logger.info(f"Claude API successful for page: {page_name} - Session: {session_id}")
        
        return jsonify({
            'success': True,
            'content': generated_content,
            'session_id': session_id,
            'page_type': 'claude-call',
            'files_processed': list(uploaded_files_data.keys()),
            'server_files_loaded': list(server_files_data.keys()),
            'form_fields_received': [k for k in form_data.keys() if k != 'claude_prompt']
        })
        
    except Exception as e:
        logger.error(f"Claude API error for page {page_name}: {e}")
        return jsonify({'error': f'Error generating content: {str(e)}'}), 500from flask import Flask, render_template, request, jsonify, session
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
import anthropic
import os
import logging
from datetime import datetime
import uuid
import io
from docx import Document
import PyPDF2

# Initialize Flask app
app = Flask(__name__)

# Simple configuration without external dependencies
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key-change-this')
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB max file size

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('logs/app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Initialize simple rate limiter (memory-based)
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["100 per hour"]
)
logger.info("Rate limiter initialized with memory storage")

# Initialize Claude client
try:
    api_key = os.environ.get('CLAUDE_API_KEY')
    if not api_key:
        raise ValueError("CLAUDE_API_KEY environment variable not set")
    if not api_key.startswith('sk-ant-'):
        raise ValueError("Invalid Claude API key format")
    
    claude_client = anthropic.Anthropic(api_key=api_key)
    logger.info("Claude API client initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize Claude client: {e}")
    logger.error(f"API key present: {bool(os.environ.get('CLAUDE_API_KEY'))}")
    logger.error(f"API key starts correctly: {str(os.environ.get('CLAUDE_API_KEY', '')).startswith('sk-ant-') if os.environ.get('CLAUDE_API_KEY') else False}")
    claude_client = None

# File processing functions
def extract_text_from_docx(file_stream):
    """Extract text from Word document"""
    try:
        doc = Document(file_stream)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Skip empty paragraphs
                text.append(paragraph.text.strip())
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(' | '.join(row_text))
        
        return '\n'.join(text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise ValueError(f"Failed to extract text from Word document: {str(e)}")

def extract_xml_from_docx(file_stream):
    """Extract full XML structure from Word document"""
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        
        # Reset stream position
        file_stream.seek(0)
        
        xml_data = {}
        
        with zipfile.ZipFile(file_stream, 'r') as docx_zip:
            # Get main document XML
            if 'word/document.xml' in docx_zip.namelist():
                document_xml = docx_zip.read('word/document.xml').decode('utf-8')
                xml_data['document'] = document_xml
            
            # Get styles XML
            if 'word/styles.xml' in docx_zip.namelist():
                styles_xml = docx_zip.read('word/styles.xml').decode('utf-8')
                xml_data['styles'] = styles_xml
            
            # Get core properties
            if 'docProps/core.xml' in docx_zip.namelist():
                core_xml = docx_zip.read('docProps/core.xml').decode('utf-8')
                xml_data['core_properties'] = core_xml
            
            # Get app properties
            if 'docProps/app.xml' in docx_zip.namelist():
                app_xml = docx_zip.read('docProps/app.xml').decode('utf-8')
                xml_data['app_properties'] = app_xml
            
            # List all files in the archive for reference
            xml_data['file_list'] = docx_zip.namelist()
        
        return xml_data
    except Exception as e:
        logger.error(f"Error extracting XML from DOCX: {e}")
        return {"error": f"Failed to extract XML structure: {str(e)}"}

def extract_text_from_pdf(file_stream):
    """Extract text from PDF document"""
    try:
        reader = PyPDF2.PdfReader(file_stream)
        text = []
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text.append(f"--- Page {page_num + 1} ---")
                    text.append(page_text.strip())
            except Exception as e:
                logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                continue
        
        if not text:
            raise ValueError("No text could be extracted from the PDF")
        
        return '\n'.join(text)
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise ValueError(f"Failed to extract text from PDF document: {str(e)}")

def extract_form_data_from_pdf(file_stream):
    """Extract form data and metadata from PDF document"""
    try:
        # Reset stream position
        file_stream.seek(0)
        reader = PyPDF2.PdfReader(file_stream)
        
        form_data = {}
        
        # Extract form fields if present
        if reader.get_form_text_fields():
            form_data['form_fields'] = reader.get_form_text_fields()
        
        # Extract metadata
        if reader.metadata:
            metadata = {}
            for key, value in reader.metadata.items():
                # Convert PyPDF2 metadata keys to readable format
                clean_key = str(key).replace('/', '').replace('\\', '')
                metadata[clean_key] = str(value) if value else None
            form_data['metadata'] = metadata
        
        # Extract document info
        form_data['document_info'] = {
            'page_count': len(reader.pages),
            'is_encrypted': reader.is_encrypted,
            'has_form_fields': bool(reader.get_form_text_fields())
        }
        
        # Extract page-level information
        page_info = []
        for page_num, page in enumerate(reader.pages):
            try:
                page_data = {
                    'page_number': page_num + 1,
                    'rotation': page.rotation if hasattr(page, 'rotation') else 0
                }
                
                # Try to get page dimensions
                if hasattr(page, 'mediabox'):
                    mediabox = page.mediabox
                    page_data['dimensions'] = {
                        'width': float(mediabox.width),
                        'height': float(mediabox.height)
                    }
                
                page_info.append(page_data)
            except Exception as e:
                logger.warning(f"Error extracting info from page {page_num + 1}: {e}")
                continue
        
        form_data['page_info'] = page_info
        
        return form_data
    except Exception as e:
        logger.error(f"Error extracting form data from PDF: {e}")
        return {"error": f"Failed to extract PDF form data: {str(e)}"}

def process_uploaded_file(file):
    """Process uploaded file and extract both content and structure"""
    if not file or not file.filename:
        return None
    
    filename = file.filename.lower()
    file_content = file.read()
    
    if not file_content:
        raise ValueError("File is empty")
    
    file_stream = io.BytesIO(file_content)
    result = {}
    
    if filename.endswith('.docx'):
        # Extract both text content and XML structure
        file_stream.seek(0)
        result['text_content'] = extract_text_from_docx(file_stream)
        
        file_stream.seek(0)
        result['xml_structure'] = extract_xml_from_docx(file_stream)
        
        result['file_type'] = 'docx'
        
    elif filename.endswith('.pdf'):
        # Extract both text content and form data
        file_stream.seek(0)
        result['text_content'] = extract_text_from_pdf(file_stream)
        
        file_stream.seek(0)
        result['form_data'] = extract_form_data_from_pdf(file_stream)
        
        result['file_type'] = 'pdf'
        
    elif filename.endswith('.txt'):
        # For text files, just get content
        result['text_content'] = file_stream.read().decode('utf-8')
        result['file_type'] = 'txt'
        
    else:
        raise ValueError(f"Unsupported file type. Please upload .docx, .pdf, or .txt files only.")
    
    return result

def process_server_file(file_path):
    """Process server-stored file and extract both content and structure"""
    if not os.path.exists(file_path):
        return None
    
    filename = file_path.lower()
    result = {}
    
    try:
        if filename.endswith('.docx'):
            with open(file_path, 'rb') as f:
                file_stream = io.BytesIO(f.read())
                
                file_stream.seek(0)
                result['text_content'] = extract_text_from_docx(file_stream)
                
                file_stream.seek(0)
                result['xml_structure'] = extract_xml_from_docx(file_stream)
                
                result['file_type'] = 'docx'
                
        elif filename.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                file_stream = io.BytesIO(f.read())
                
                file_stream.seek(0)
                result['text_content'] = extract_text_from_pdf(file_stream)
                
                file_stream.seek(0)
                result['form_data'] = extract_form_data_from_pdf(file_stream)
                
                result['file_type'] = 'pdf'
                
        elif filename.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                result['text_content'] = f.read()
                result['file_type'] = 'txt'
        else:
            logger.warning(f"Unsupported server file type: {file_path}")
            return None
            
        return result
    except Exception as e:
        logger.error(f"Error processing server file {file_path}: {e}")
        return None

def load_server_files(page_name):
    """Load all files from the server directory for this page"""
    server_files = {}
    server_dir = f"/app/server_files/{page_name}"
    
    logger.info(f"Looking for server files in: {server_dir}")
    
    if not os.path.exists(server_dir):
        logger.warning(f"No server files directory found for page: {page_name} at path: {server_dir}")
        return server_files
    
    try:
        all_files = os.listdir(server_dir)
        logger.info(f"Files found in directory: {all_files}")
        
        for filename in all_files:
            file_path = os.path.join(server_dir, filename)
            logger.info(f"Processing file: {file_path}")
            
            if os.path.isfile(file_path):
                logger.info(f"File {filename} is a regular file, attempting to process...")
                content = process_server_file(file_path)
                if content:
                    # Use filename without extension as key
                    file_key = os.path.splitext(filename)[0].replace('_', ' ').replace('-', ' ')
                    server_files[file_key] = content
                    logger.info(f"Successfully loaded server file: {filename} as key: {file_key}")
                else:
                    logger.warning(f"Failed to extract content from file: {filename}")
            else:
                logger.info(f"Skipping {filename} - not a regular file")
        
        logger.info(f"Total server files loaded for page {page_name}: {len(server_files)}")
        if server_files:
            logger.info(f"Server file keys: {list(server_files.keys())}")
        
    except Exception as e:
        logger.error(f"Error loading server files for page {page_name}: {e}")
    
    return server_files

def validate_file(file):
    """Validate uploaded file"""
    if not file or not file.filename:
        return False, "No file provided"
    
    filename = file.filename.lower()
    if not (filename.endswith('.docx') or filename.endswith('.pdf')):
        return False, "Only .docx and .pdf files are supported"
    
    # Check file size (this is also enforced by Flask config)
    file.seek(0, 2)  # Seek to end
    size = file.tell()
    file.seek(0)  # Reset to beginning
    
    if size > 10 * 1024 * 1024:  # 10MB
        return False, "File size must be less than 10MB"
    
    if size == 0:
        return False, "File is empty"
    
    return True, "File is valid"

def build_claude_prompt(form_data, uploaded_files_text, server_files_text):
    """Build Claude prompt from form data, uploaded files, and server files"""
    
    # Get the custom prompt from the form
    custom_prompt = form_data.get('claude_prompt', '').strip()
    if not custom_prompt:
        return None, "No Claude prompt provided"
    
    # Build the context sections
    context_sections = []
    
    # Add server files first (background/examples)
    if server_files_text:
        context_sections.append("=== SERVER REFERENCE FILES ===")
        for file_name, content in server_files_text.items():
            context_sections.append(f"\n--- {file_name.upper()} ---")
            context_sections.append(content.strip())
    
    # Add form data (excluding the prompt itself)
    form_context = {}
    for key, value in form_data.items():
        if key != 'claude_prompt' and value.strip():
            form_context[key] = value.strip()
    
    if form_context:
        context_sections.append("\n=== FORM DATA ===")
        for key, value in form_context.items():
            context_sections.append(f"{key.replace('_', ' ').title()}: {value}")
    
    # Add uploaded file contents
    if uploaded_files_text:
        for file_type, content in uploaded_files_text.items():
            if content.strip():
                context_sections.append(f"\n=== UPLOADED {file_type.upper()} CONTENT ===")
                context_sections.append(content.strip())
    
    # Combine everything
    if context_sections:
        full_prompt = "\n".join(context_sections) + "\n\n" + custom_prompt
    else:
        full_prompt = custom_prompt
    
    return full_prompt, None

# Utility function to generate session ID
def get_session_id():
    if 'session_id' not in session:
        session['session_id'] = str(uuid.uuid4())
    return session['session_id']

# Home page
@app.route('/')
def home():
    session_id = get_session_id()
    logger.info(f"Home page accessed - Session: {session_id}")
    return render_template('home.html')

# Health check endpoint
@app.route('/health')
def health():
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.utcnow().isoformat(),
        'claude_available': claude_client is not None
    })

# Favicon route to prevent errors
@app.route('/favicon.ico')
def favicon():
    return '', 204  # No content response

# Generic page route
@app.route('/<page_name>')
def generic_page(page_name):
    """Serve any page that has a corresponding template"""
    session_id = get_session_id()
    logger.info(f"Page accessed: {page_name} - Session: {session_id}")
    
    # Convert URL format to template format (e.g., chairs-promotion-letter -> chairs_promotion_letter.html)
    template_name = page_name.replace('-', '_') + '.html'
    
    # Load server files for this page to show on the page
    try:
        server_files_info = get_server_files_info(page_name)
        logger.info(f"Server files info loaded for {page_name}: {len(server_files_info)} files")
    except Exception as e:
        logger.error(f"Error loading server files info for {page_name}: {e}")
        server_files_info = []
    
    try:
        return render_template(template_name, page_name=page_name, server_files_info=server_files_info)
    except Exception as e:
        logger.error(f"Template error for {template_name}: {e}")
        return render_template('404.html'), 404

def get_server_files_info(page_name):
    """Get information about available server files for display"""
    server_files_info = []
    server_dir = f"/app/server_files/{page_name}"
    
    logger.info(f"Getting server files info for page: {page_name} from directory: {server_dir}")
    
    if not os.path.exists(server_dir):
        logger.info(f"Server files directory does not exist: {server_dir}")
        return server_files_info
    
    try:
        files_list = os.listdir(server_dir)
        logger.info(f"Found {len(files_list)} files in {server_dir}: {files_list}")
        
        for filename in files_list:
            try:
                file_path = os.path.join(server_dir, filename)
                logger.info(f"Processing file: {file_path}")
                
                if not os.path.isfile(file_path):
                    logger.info(f"Skipping {filename} - not a regular file")
                    continue
                
                # Get file info
                file_stat = os.stat(file_path)
                file_size = file_stat.st_size
                
                # Format file size
                if file_size < 1024:
                    size_str = f"{file_size} B"
                elif file_size < 1024 * 1024:
                    size_str = f"{file_size / 1024:.1f} KB"
                else:
                    size_str = f"{file_size / (1024 * 1024):.1f} MB"
                
                # Get file type
                file_ext = os.path.splitext(filename)[1].lower()
                if file_ext == '.docx':
                    file_type = 'Word Document'
                elif file_ext == '.pdf':
                    file_type = 'PDF Document'
                elif file_ext == '.txt':
                    file_type = 'Text File'
                else:
                    file_type = 'Unknown'
                
                # Create display name
                display_name = os.path.splitext(filename)[0].replace('_', ' ').replace('-', ' ').title()
                
                file_info = {
                    'filename': filename,
                    'display_name': display_name,
                    'file_type': file_type,
                    'size': size_str,
                    'supported': file_ext in ['.docx', '.pdf', '.txt']
                }
                
                server_files_info.append(file_info)
                logger.info(f"Added file info: {file_info}")
                
            except Exception as file_error:
                logger.error(f"Error processing file {filename}: {file_error}")
                continue
        
        # Sort by filename
        server_files_info.sort(key=lambda x: x['filename'])
        logger.info(f"Successfully processed {len(server_files_info)} files for page {page_name}")
        
    except Exception as e:
        logger.error(f"Error listing files in directory {server_dir}: {e}")
        return []
    
    return server_files_info

# Generic API endpoint
@app.route('/api/<page_name>', methods=['POST'])
@limiter.limit("5 per minute")
def generic_api(page_name):
    """Generic API endpoint that handles any page"""
    try:
        # Handle file uploads - process all uploaded files
        uploaded_files_data = {}
        
        for field_name in request.files:
            file = request.files[field_name]
            if file and file.filename:
                is_valid, message = validate_file(file)
                if not is_valid:
                    return jsonify({'error': f'{field_name} error: {message}'}), 400
                
                try:
                    file_data = process_uploaded_file(file)
                    if file_data:
                        # Clean up field name for context
                        clean_field_name = field_name.replace('_file', '').replace('_', ' ')
                        uploaded_files_data[clean_field_name] = file_data
                        logger.info(f"File processed: {file.filename} for field {field_name}")
                except Exception as e:
                    return jsonify({'error': f'Error processing {field_name}: {str(e)}'}), 400
        
        # Load server files for this page
        server_files_data = load_server_files(page_name)
        
        # Handle form data
        form_data = request.form.to_dict()
        session_id = get_session_id()
        
        logger.info(f"API called for page: {page_name} - Session: {session_id}")
        
        # Check page type based on form data
        page_type = form_data.get('page_type', 'claude-call')
        
        if page_type == 'no-call':
            # Handle no-call pages (no Claude API, just return organized data)
            return handle_no_call_page(page_name, form_data, uploaded_files_data, server_files_data, session_id)
        elif page_type == 'claude-call':
            # Handle Claude API pages (existing functionality)
            return handle_claude_call_page(page_name, form_data, uploaded_files_data, server_files_data, session_id)
        else:
            return jsonify({'error': f'Unknown page type: {page_type}'}), 400
    
    except Exception as e:
        logger.error(f"Error in API for page {page_name}: {e}")
        return jsonify({'error': 'Internal server error'}), 500

def handle_no_call_page(page_name, form_data, uploaded_files_text, server_files_text, session_id):
    """Handle no-call pages that organize and display information without API calls"""
    try:
        logger.info(f"Processing no-call page: {page_name}")
        
        # Build organized content from available data
        content_sections = []
        
        # Add page header
        page_title = form_data.get('page_title', page_name.replace('-', ' ').title())
        content_sections.append(f"# {page_title}")
        content_sections.append("")
        
        # Add server reference files
        if server_files_text:
            content_sections.append("## Reference Materials")
            content_sections.append("")
            for file_name, content in server_files_text.items():
                content_sections.append(f"### {file_name.title()}")
                content_sections.append("")
                # Add truncated preview of content
                preview = content[:500] + "..." if len(content) > 500 else content
                content_sections.append(preview)
                content_sections.append("")
        
        # Add form data
        if form_data:
            content_sections.append("## Submitted Information")
            content_sections.append("")
            for key, value in form_data.items():
                if key not in ['page_type', 'page_title'] and value.strip():
                    display_key = key.replace('_', ' ').title()
                    content_sections.append(f"**{display_key}:** {value}")
            content_sections.append("")
        
        # Add uploaded files
        if uploaded_files_text:
            content_sections.append("## Uploaded Documents")
            content_sections.append("")
            for file_type, content in uploaded_files_text.items():
                content_sections.append(f"### {file_type.title()}")
                content_sections.append("")
                # Add truncated preview
                preview = content[:500] + "..." if len(content) > 500 else content
                content_sections.append(preview)
                content_sections.append("")
        
        # Add summary
        total_items = len(server_files_text) + len(uploaded_files_text) + len([k for k in form_data.keys() if k not in ['page_type', 'page_title'] and form_data[k].strip()])
        content_sections.append("## Summary")
        content_sections.append("")
        content_sections.append(f"- **Server Reference Files:** {len(server_files_text)}")
        content_sections.append(f"- **Uploaded Documents:** {len(uploaded_files_text)}")
        content_sections.append(f"- **Form Fields Completed:** {len([k for k in form_data.keys() if k not in ['page_type', 'page_title'] and form_data[k].strip()])}")
        content_sections.append(f"- **Total Items Processed:** {total_items}")
        
        generated_content = "\n".join(content_sections)
        
        logger.info(f"No-call page processed successfully for: {page_name}")
        
        return jsonify({
            'success': True,
            'content': generated_content,
            'session_id': session_id,
            'page_type': 'no-call',
            'files_processed': list(uploaded_files_text.keys()),
            'server_files_loaded': list(server_files_text.keys()),
            'form_fields_received': [k for k in form_data.keys() if k not in ['page_type', 'page_title']]
        })
        
    except Exception as e:
        logger.error(f"Error in no-call page handler: {e}")
        return jsonify({'error': 'Error processing no-call page'}), 500

def handle_claude_call_page(page_name, form_data, uploaded_files_text, server_files_text, session_id):
    """Handle Claude API pages (existing functionality)"""
    if not claude_client:
        return jsonify({'error': 'Claude API not available'}), 503
    
    try:
        # Build Claude prompt from form data, uploaded files, and server files
        claude_prompt, error = build_claude_prompt(form_data, uploaded_files_text, server_files_text)
        if error:
            return jsonify({'error': error}), 400
        
        # Call Claude API
        message = claude_client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=4000,
            messages=[
                {"role": "user", "content": claude_prompt}
            ]
        )
        
        generated_content = message.content[0].text
        
        logger.info(f"Claude API successful for page: {page_name} - Session: {session_id}")
        
        return jsonify({
            'success': True,
            'content': generated_content,
            'session_id': session_id,
            'page_type': 'claude-call',
            'files_processed': list(uploaded_files_text.keys()),
            'server_files_loaded': list(server_files_text.keys()),
            'form_fields_received': [k for k in form_data.keys() if k != 'claude_prompt']
        })
        
    except Exception as e:
        logger.error(f"Claude API error for page {page_name}: {e}")
        return jsonify({'error': f'Error generating content: {str(e)}'}), 500

# Error handlers
@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'File too large. Maximum size is 10MB.'}), 413

@app.errorhandler(429)
def ratelimit_handler(e):
    return jsonify({'error': 'Rate limit exceeded. Please try again later.'}), 429

@app.errorhandler(404)
def not_found(e):
    try:
        return render_template('404.html'), 404
    except:
        # Fallback if 404.html template fails
        return '''
        <html><body>
        <h1>404 Page Not Found</h1>
        <p>The page you're looking for doesn't exist.</p>
        <a href="/">Go Home</a>
        </body></html>
        ''', 404

@app.errorhandler(500)
def internal_error(e):
    logger.error(f"Internal server error: {e}")
    try:
        return render_template('500.html'), 500
    except:
        # Fallback if 500.html template fails
        return '''
        <html><body>
        <h1>500 Internal Server Error</h1>
        <p>The server encountered an internal error.</p>
        <a href="/">Go Home</a>
        </body></html>
        ''', 500

if __name__ == '__main__':
    # Create logs directory if it doesn't exist
    os.makedirs('logs', exist_ok=True)
    
    # Run the application
    app.run(
        host='0.0.0.0',
        port=5000,
        debug=os.environ.get('FLASK_DEBUG', 'false').lower() == 'true'
    )
