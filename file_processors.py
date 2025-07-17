"""
File processing functions for Word documents, PDFs, and text files
"""

import os
import io
import logging
import zipfile
import json
from docx import Document
import PyPDF2

logger = logging.getLogger(__name__)

def extract_text_from_docx(file_stream):
    """Extract text from Word document"""
    try:
        doc = Document(file_stream)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Skip empty paragraphs
                text.append(paragraph.text.strip())
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(' | '.join(row_text))
        
        return '\n'.join(text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise ValueError(f"Failed to extract text from Word document: {str(e)}")

def extract_xml_from_docx(file_stream):
    """Extract full XML structure from Word document"""
    try:
        # Reset stream position
        file_stream.seek(0)
        
        xml_data = {}
        
        with zipfile.ZipFile(file_stream, 'r') as docx_zip:
            # Get main document XML
            if 'word/document.xml' in docx_zip.namelist():
                document_xml = docx_zip.read('word/document.xml').decode('utf-8')
                xml_data['document'] = document_xml
            
            # Get styles XML
            if 'word/styles.xml' in docx_zip.namelist():
                styles_xml = docx_zip.read('word/styles.xml').decode('utf-8')
                xml_data['styles'] = styles_xml
            
            # Get core properties
            if 'docProps/core.xml' in docx_zip.namelist():
                core_xml = docx_zip.read('docProps/core.xml').decode('utf-8')
                xml_data['core_properties'] = core_xml
            
            # Get app properties
            if 'docProps/app.xml' in docx_zip.namelist():
                app_xml = docx_zip.read('docProps/app.xml').decode('utf-8')
                xml_data['app_properties'] = app_xml
            
            # List all files in the archive for reference
            xml_data['file_list'] = docx_zip.namelist()
        
        return xml_data
    except Exception as e:
        logger.error(f"Error extracting XML from DOCX: {e}")
        return {"error": f"Failed to extract XML structure: {str(e)}"}

def extract_text_from_pdf(file_stream):
    """Extract text from PDF document"""
    try:
        reader = PyPDF2.PdfReader(file_stream)
        text = []
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text.append(f"--- Page {page_num + 1} ---")
                    text.append(page_text.strip())
            except Exception as e:
                logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                continue
        
        if not text:
            raise ValueError("No text could be extracted from the PDF")
        
        return '\n'.join(text)
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise ValueError(f"Failed to extract text from PDF document: {str(e)}")

def extract_form_data_from_pdf(file_stream):
    """Extract form data and metadata from PDF document"""
    try:
        # Reset stream position
        file_stream.seek(0)
        reader = PyPDF2.PdfReader(file_stream)
        
        form_data = {}
        
        # Extract form fields if present
        if reader.get_form_text_fields():
            form_data['form_fields'] = reader.get_form_text_fields()
        
        # Extract metadata
        if reader.metadata:
            metadata = {}
            for key, value in reader.metadata.items():
                # Convert PyPDF2 metadata keys to readable format
                clean_key = str(key).replace('/', '').replace('\\', '')
                metadata[clean_key] = str(value) if value else None
            form_data['metadata'] = metadata
        
        # Extract document info
        form_data['document_info'] = {
            'page_count': len(reader.pages),
            'is_encrypted': reader.is_encrypted,
            'has_form_fields': bool(reader.get_form_text_fields())
        }
        
        # Extract page-level information
        page_info = []
        for page_num, page in enumerate(reader.pages):
            try:
                page_data = {
                    'page_number': page_num + 1,
                    'rotation': page.rotation if hasattr(page, 'rotation') else 0
                }
                
                # Try to get page dimensions
                if hasattr(page, 'mediabox'):
                    mediabox = page.mediabox
                    page_data['dimensions'] = {
                        'width': float(mediabox.width),
                        'height': float(mediabox.height)
                    }
                
                page_info.append(page_data)
            except Exception as e:
                logger.warning(f"Error extracting info from page {page_num + 1}: {e}")
                continue
        
        form_data['page_info'] = page_info
        
        return form_data
    except Exception as e:
        logger.error(f"Error extracting form data from PDF: {e}")
        return {"error": f"Failed to extract PDF form data: {str(e)}"}

def process_uploaded_file(file):
    """Process uploaded file and extract both content and structure"""
    if not file or not file.filename:
        return None
    
    filename = file.filename.lower()
    file_content = file.read()
    
    if not file_content:
        raise ValueError("File is empty")
    
    file_stream = io.BytesIO(file_content)
    result = {}
    
    if filename.endswith('.docx'):
        # Extract both text content and XML structure
        file_stream.seek(0)
        result['text_content'] = extract_text_from_docx(file_stream)
        
        file_stream.seek(0)
        result['xml_structure'] = extract_xml_from_docx(file_stream)
        
        result['file_type'] = 'docx'
        
    elif filename.endswith('.pdf'):
        # Extract both text content and form data
        file_stream.seek(0)
        result['text_content'] = extract_text_from_pdf(file_stream)
        
        file_stream.seek(0)
        result['form_data'] = extract_form_data_from_pdf(file_stream)
        
        result['file_type'] = 'pdf'
        
    elif filename.endswith('.txt'):
        # For text files, just get content
        result['text_content'] = file_stream.read().decode('utf-8')
        result['file_type'] = 'txt'
        
    else:
        raise ValueError(f"Unsupported file type. Please upload .docx, .pdf, or .txt files only.")
    
    return result

def process_server_file(file_path):
    """Process server-stored file and extract both content and structure"""
    if not os.path.exists(file_path):
        return None
    
    filename = file_path.lower()
    result = {}
    
    try:
        if filename.endswith('.docx'):
            with open(file_path, 'rb') as f:
                file_stream = io.BytesIO(f.read())
                
                file_stream.seek(0)
                result['text_content'] = extract_text_from_docx(file_stream)
                
                file_stream.seek(0)
                result['xml_structure'] = extract_xml_from_docx(file_stream)
                
                result['file_type'] = 'docx'
                
        elif filename.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                file_stream = io.BytesIO(f.read())
                
                file_stream.seek(0)
                result['text_content'] = extract_text_from_pdf(file_stream)
                
                file_stream.seek(0)
                result['form_data'] = extract_form_data_from_pdf(file_stream)
                
                result['file_type'] = 'pdf'
                
        elif filename.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                result['text_content'] = f.read()
                result['file_type'] = 'txt'
        else:
            logger.warning(f"Unsupported server file type: {file_path}")
            return None
            
        return result
    except Exception as e:
        logger.error(f"Error processing server file {file_path}: {e}")
        return None

def validate_file(file):
    """Validate uploaded file"""
    if not file or not file.filename:
        return False, "No file provided"
    
    filename = file.filename.lower()
    if not (filename.endswith('.docx') or filename.endswith('.pdf') or filename.endswith('.txt')):
        return False, "Only .docx, .pdf, and .txt files are supported"
    
    # Check file size (this is also enforced by Flask config)
    file.seek(0, 2)  # Seek to end
    size = file.tell()
    file.seek(0)  # Reset to beginning
    
    if size > 10 * 1024 * 1024:  # 10MB
        return False, "File size must be less than 10MB"
    
    if size == 0:
        return False, "File is empty"
    
    return True, "File is valid"

def load_server_files(page_name, directories=None):
    """Load all files from the server directory/directories for this page
    
    Args:
        page_name: The name of the current page
        directories: Optional list of directories to load from. If None, loads only from page_name directory
    
    Returns:
        dict: Server files with keys prefixed by directory name if loading from multiple directories
    """
    server_files = {}
    
    # Default to just the page's own directory if not specified
    if directories is None:
        directories = [page_name]
    
    # Single directory mode - no prefixing
    if len(directories) == 1:
        server_dir = f"/app/server_files/{directories[0]}"
        
        logger.info(f"Looking for server files in: {server_dir}")
        
        if not os.path.exists(server_dir):
            logger.warning(f"No server files directory found for page: {directories[0]} at path: {server_dir}")
            return server_files
        
        try:
            all_files = os.listdir(server_dir)
            logger.info(f"Files found in directory: {all_files}")
            
            for filename in all_files:
                file_path = os.path.join(server_dir, filename)
                logger.info(f"Processing file: {file_path}")
                
                if os.path.isfile(file_path):
                    logger.info(f"File {filename} is a regular file, attempting to process...")
                    file_data = process_server_file(file_path)
                    if file_data:
                        # Use filename without extension as key
                        file_key = os.path.splitext(filename)[0].replace('_', ' ').replace('-', ' ')
                        server_files[file_key] = file_data
                        logger.info(f"Successfully loaded server file: {filename} as key: {file_key}")
                    else:
                        logger.warning(f"Failed to extract content from file: {filename}")
                else:
                    logger.info(f"Skipping {filename} - not a regular file")
            
            logger.info(f"Total server files loaded for page {page_name}: {len(server_files)}")
            if server_files:
                logger.info(f"Server file keys: {list(server_files.keys())}")
        
        except Exception as e:
            logger.error(f"Error loading server files for page {page_name}: {e}")
    
    else:
        # Multiple directories mode - prefix with directory name
        logger.info(f"Loading server files from multiple directories for page {page_name}: {directories}")
        
        for directory in directories:
            server_dir = f"/app/server_files/{directory}"
            
            if not os.path.exists(server_dir):
                logger.info(f"Directory {server_dir} does not exist, skipping")
                continue
            
            try:
                all_files = os.listdir(server_dir)
                logger.info(f"Files found in {directory}: {all_files}")
                
                for filename in all_files:
                    file_path = os.path.join(server_dir, filename)
                    
                    if os.path.isfile(file_path):
                        file_data = process_server_file(file_path)
                        if file_data:
                            # Create key with directory prefix
                            base_key = os.path.splitext(filename)[0].replace('_', ' ').replace('-', ' ')
                            # Only prefix if not loading from the page's own directory
                            if directory == page_name:
                                file_key = base_key
                            else:
                                file_key = f"{directory} - {base_key}"
                            
                            # Handle potential key conflicts
                            if file_key in server_files:
                                logger.warning(f"Key conflict for {file_key}, keeping first occurrence")
                                continue
                            
                            server_files[file_key] = file_data
                            logger.info(f"Loaded {filename} from {directory} as '{file_key}'")
                        else:
                            logger.warning(f"Failed to extract content from {directory}/{filename}")
                
            except Exception as e:
                logger.error(f"Error loading files from directory {directory}: {e}")
                continue
        
        logger.info(f"Total server files loaded from all directories: {len(server_files)}")
        if server_files:
            logger.info(f"All server file keys: {list(server_files.keys())}")
    
    return server_files
